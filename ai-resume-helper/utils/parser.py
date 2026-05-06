"""
文件解析器模块
支持 PDF 和 DOCX 格式的简历文件解析
"""

import pdfplumber
from docx import Document
from io import BytesIO


def parse_pdf(file_stream) -> tuple:
    """
    使用 pdfplumber 解析 PDF 文件

    Args:
        file_stream: 文件流对象

    Returns:
        tuple: (提取的文本内容, 错误信息)
    """
    try:
        text_parts = []
        with pdfplumber.open(file_stream) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

        full_text = "\n".join(text_parts).strip()

        if not full_text:
            return None, "PDF 文件中未提取到任何文本内容"

        return full_text, None

    except Exception as e:
        return None, f"PDF 解析失败: {str(e)}"


def parse_docx(file_stream) -> tuple:
    """
    使用 python-docx 解析 DOCX 文件

    Args:
        file_stream: 文件流对象

    Returns:
        tuple: (提取的文本内容, 错误信息)
    """
    try:
        doc = Document(file_stream)
        text_parts = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                text_parts.append(text)

        # 提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_parts.append(" | ".join(row_text))

        full_text = "\n".join(text_parts).strip()

        if not full_text:
            return None, "DOCX 文件中未提取到任何文本内容"

        return full_text, None

    except Exception as e:
        return None, f"DOCX 解析失败: {str(e)}"


def parse_resume(file_storage) -> tuple:
    """
    统一的简历解析入口

    Args:
        file_storage: Flask request.files 中的文件对象

    Returns:
        tuple: (提取的文本内容, 错误信息)
    """
    if not file_storage:
        return None, "未上传文件"

    filename = file_storage.filename.lower()

    if filename.endswith('.pdf'):
        return parse_pdf(file_storage)

    elif filename.endswith('.docx'):
        return parse_docx(file_storage)

    else:
        return None, "不支持的文件格式，仅支持 PDF 和 DOCX"
